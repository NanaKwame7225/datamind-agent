"""
DataMind Agent — Document Export Service
Generates professional Word (.docx) and PDF reports.
Times New Roman, 12pt body, 14pt headings, bulleted paragraphs.
"""
from __future__ import annotations
import io, logging
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

# ── AUTHOR SIGNATURE — appears on every generated document ────────────────────
SIGNATURE = {
    "name": "Nana Kwame Asomani-Appah",
    "phone1": "+233 55 702 3768",
    "phone2": "+233 53 835 0574",
    "linkedin": "https://www.linkedin.com/in/nana-kwame-asomani",
    "org": "NkaySolutions",
    "location": "Accra, Ghana",
}


class DocumentExportService:
    """
    Exports analysis reports as professional documents.
    Word: python-docx. PDF: reportlab.
    Times New Roman throughout, 12pt body, 14pt headings.
    """

    # ── WORD EXPORT ───────────────────────────────────────────────────────────

    def build_word_report(self, result: dict, finance: dict = None, chart_images: list = None,
                          doc_title: str = None, doc_subtitle: str = None,
                          line_spacing: float = 1.5) -> bytes:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Base style — Times New Roman 12pt
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        # Justified text with the requested line spacing throughout.
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
            style.paragraph_format.alignment = _AL.JUSTIFY
            style.paragraph_format.line_spacing = float(line_spacing or 1.5)
        except Exception:
            pass

        industry = result.get("industry", "General").replace("_", " ").title()
        query = result.get("query", "Data Analysis")
        date_str = datetime.now(timezone.utc).strftime("%d %B %Y")

        # ── TITLE (user-supplied; no auto branding header) ──
        if doc_title:
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run(str(doc_title))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.font.bold = True
        if doc_subtitle:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run(str(doc_subtitle))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.italic = True
        if doc_title or doc_subtitle:
            doc.add_paragraph()

        # ── QUERY ──
        self._add_heading(doc, "Analysis Question")
        self._add_body(doc, query)

        # ── METRICS ──
        metrics = result.get("metrics", [])
        if metrics:
            self._add_heading(doc, "Key Metrics")
            for m in metrics:
                change = f" ({m.get('change_pct','')}% vs previous)" if m.get('change_pct') is not None else ""
                self._add_bullet(doc, f"{m.get('label','')}: {m.get('value','')}{change}")

        # ── AI NARRATIVE ──
        narrative = result.get("narrative", "")
        if narrative:
            self._add_heading(doc, "Analysis & Recommendations")
            self._render_narrative_word(doc, narrative)

        # ── FINDINGS ──
        insights = result.get("insights", [])
        if insights:
            self._add_heading(doc, "Key Findings")
            for i in insights:
                sev = i.get("severity", "info").upper()
                conf = f" (Confidence: {int(i.get('confidence',0)*100)}%)" if i.get('confidence') else ""
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"[{sev}] {i.get('title','')}{conf}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
                body_p = doc.add_paragraph()
                body_p.paragraph_format.left_indent = Inches(0.5)
                run = body_p.add_run(i.get('body',''))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                if i.get('source'):
                    src_p = doc.add_paragraph()
                    src_p.paragraph_format.left_indent = Inches(0.5)
                    run = src_p.add_run(f"Method: {i.get('source')}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # ── CHART IMAGES ──
        if chart_images:
            self._add_charts_word(doc, chart_images)

        # ── FINANCE MODULES ──
        if finance:
            self._add_finance_word(doc, finance)

        # ── SIGNATURE BLOCK ──
        self._add_signature_word(doc)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def _add_signature_word(self, doc):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc.add_paragraph()

        # Horizontal rule
        hr = doc.add_paragraph()
        pPr = hr._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'cccccc')
        pbdr.append(bottom)
        pPr.append(pbdr)

        # Prepared by
        p = doc.add_paragraph()
        r = p.add_run("Prepared by")
        r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        p = doc.add_paragraph()
        r = p.add_run(SIGNATURE["name"])
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
        r.font.color.rgb = RGBColor(0x0a, 0x4d, 0x4a)

        for line in [
            f"Telephone: {SIGNATURE['phone1']}  |  {SIGNATURE['phone2']}",
            f"LinkedIn: {SIGNATURE['linkedin']}",
        ]:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # (Platform footer intentionally omitted.)

    def _add_heading(self, doc, text):
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0a, 0x4d, 0x4a)
        # bottom border
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), '0a4d4a')
        pbdr.append(bottom)
        pPr.append(pbdr)

    def _add_body(self, doc, text):
        from docx.shared import Pt
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def _add_bullet(self, doc, text):
        from docx.shared import Pt
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def _add_word_table(self, doc, header, body):
        """A real Word table: header band, banded rows, right-aligned numbers."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if not header:
            return
        table = doc.add_table(rows=1, cols=len(header))
        try:
            table.style = 'Light Grid Accent 1'
        except Exception:
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
        # Header row
        hdr = table.rows[0].cells
        for k, h in enumerate(header):
            hdr[k].text = ''
            p = hdr[k].paragraphs[0]
            run = p.add_run(str(h))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0a, 0x4d, 0x4a)
        # Body
        def is_num(v):
            try:
                float(str(v).replace(',', '').replace('%', ''))
                return True
            except Exception:
                return False
        for row in body:
            cells = table.add_row().cells
            for k in range(len(header)):
                v = row[k] if k < len(row) else ''
                cells[k].text = ''
                p = cells[k].paragraphs[0]
                if is_num(v):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(str(v))
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        doc.add_paragraph()

    def _render_narrative_word(self, doc, narrative):
        from docx.shared import Pt, RGBColor
        # Clean JSON leakage
        import re
        narrative = re.sub(r'^\s*\{"narrative"\s*:\s*"', '', narrative)
        narrative = re.sub(r'",\s*"metrics"\s*:\s*\[.*$', '', narrative, flags=re.DOTALL)
        narrative = narrative.replace('\\n', '\n').replace('\\"', '"')

        # Single pass so everything renders in document order. Markdown tables
        # become REAL tables — as plain paragraphs they read as a wall of pipes.
        lines = narrative.split('\n')
        i = -1
        while i + 1 < len(lines):
            i += 1
            line = lines[i].strip()
            if not line:
                continue
            # ── Markdown table: a header row followed by a |---|---| separator
            if (re.match(r'^\|.*\|$', line) and i + 1 < len(lines)
                    and re.match(r'^\|[\s:\-\|]+\|$', lines[i + 1].strip())):
                header = [c.strip() for c in line.strip().strip('|').split('|')]
                j = i + 2
                body = []
                while j < len(lines) and re.match(r'^\|.*\|$', lines[j].strip()):
                    body.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                    j += 1
                self._add_word_table(doc, header, body)
                i = j - 1
                continue
            # Markdown headings
            if line.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                run = p.add_run(line[3:].strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0a, 0x4d, 0x4a)
            elif line.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(line[4:].strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.bold = True
            # Bulleted / numbered lines become bullet paragraphs
            elif line.startswith(('- ', '• ', '* ')) or re.match(r'^\d+\.\s', line):
                clean = re.sub(r'^[-•*]\s*', '', line)
                clean = re.sub(r'^\d+\.\s*', '', clean)
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', clean)
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(clean)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            else:
                # Regular paragraph with bold support
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    def _add_charts_word(self, doc, chart_images):
        import base64, io
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        self._add_heading(doc, "Dashboard Visualisations")
        for ci in chart_images[:8]:
            if not ci or not ci.get("image"):
                continue
            title = ci.get("title", "")
            if title:
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.font.bold = True
            try:
                img_bytes = base64.b64decode(ci["image"].split(",")[-1])
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
            except Exception:
                pass
            if ci.get("subtitle"):
                sp = doc.add_paragraph()
                r = sp.add_run(ci["subtitle"])
                r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.italic = True
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_finance_word(self, doc, finance):
        from docx.shared import Pt
        tax = finance.get("tax")
        acct = finance.get("accounting")
        fraud = finance.get("fraud")

        if tax and not tax.get("error"):
            self._add_heading(doc, "Tax Analysis")
            for m in tax.get("metrics", []):
                self._add_bullet(doc, f"{m.get('label','')}: {m.get('value','')} (benchmark: {m.get('benchmark','')})")
            for f in tax.get("findings", [])[:5]:
                self._add_bullet(doc, f"{f.get('title','')} — {f.get('body','')}")

        if acct and not acct.get("error"):
            self._add_heading(doc, "Accounting Analysis")
            self._add_body(doc, f"Balance Sheet Health Score: {acct.get('health_score','?')}/100")
            for m in acct.get("metrics", []):
                self._add_bullet(doc, f"{m.get('label','')}: {m.get('value','')}")

        if fraud and not fraud.get("error"):
            self._add_heading(doc, "Fraud Detection")
            self._add_body(doc, f"Risk Score: {fraud.get('risk_score','?')}/100 ({fraud.get('risk_level','?')} risk)")
            for f in fraud.get("findings", [])[:5]:
                if f.get("severity") in ["critical", "warning"]:
                    self._add_bullet(doc, f"{f.get('title','')} — {f.get('body','')}")

        actions = finance.get("priority_actions", [])
        if actions:
            self._add_heading(doc, "Priority Actions")
            for a in actions:
                self._add_bullet(doc, f"[{a.get('module','')}] {a.get('action','')} — {a.get('reason','')}")

    # ── PDF EXPORT ────────────────────────────────────────────────────────────

    def build_pdf_report(self, result: dict, finance: dict = None, chart_images: list = None,
                         doc_title: str = None, doc_subtitle: str = None,
                         line_spacing: float = 1.5) -> bytes:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        import re

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
                                topMargin=0.8*inch, bottomMargin=0.8*inch,
                                leftMargin=1*inch, rightMargin=1*inch)

        # Styles — Times New Roman
        title_style = ParagraphStyle('Title', fontName='Times-Bold', fontSize=20,
                                     textColor=HexColor('#0a4d4a'), alignment=TA_CENTER, spaceAfter=4)
        subtitle_style = ParagraphStyle('Subtitle', fontName='Times-Bold', fontSize=14,
                                        alignment=TA_CENTER, spaceAfter=4)
        meta_style = ParagraphStyle('Meta', fontName='Times-Roman', fontSize=10,
                                    textColor=HexColor('#666666'), alignment=TA_CENTER, spaceAfter=16)
        heading_style = ParagraphStyle('Heading', fontName='Times-Bold', fontSize=14,
                                       textColor=HexColor('#0a4d4a'), spaceBefore=14, spaceAfter=6,
                                       borderWidth=0, borderColor=HexColor('#0a4d4a'), borderPadding=2)
        subhead_style = ParagraphStyle('Subhead', fontName='Times-Bold', fontSize=12,
                                       spaceBefore=8, spaceAfter=3)
        _ls = float(line_spacing or 1.5)
        body_style = ParagraphStyle('Body', fontName='Times-Roman', fontSize=12,
                                    spaceAfter=6, leading=12 * _ls,
                                    alignment=4)  # 4 = TA_JUSTIFY
        bullet_style = ParagraphStyle('Bullet', fontName='Times-Roman', fontSize=12,
                                      leftIndent=20, spaceAfter=4, leading=12 * _ls,
                                      alignment=4)
        src_style = ParagraphStyle('Source', fontName='Times-Italic', fontSize=10,
                                   textColor=HexColor('#666666'), leftIndent=20, spaceAfter=6)
        footer_style = ParagraphStyle('Footer', fontName='Times-Italic', fontSize=9,
                                      textColor=HexColor('#999999'), alignment=TA_CENTER, spaceBefore=20)

        story = []
        industry = result.get("industry", "General").replace("_", " ").title()
        query = result.get("query", "Data Analysis")
        date_str = datetime.now(timezone.utc).strftime("%d %B %Y")

        # Title
        if doc_title:
            story.append(Paragraph(self._esc(str(doc_title)), title_style))
        if doc_subtitle:
            story.append(Paragraph(self._esc(str(doc_subtitle)), subtitle_style))

        # Query
        story.append(Paragraph("Analysis Question", heading_style))
        story.append(Paragraph(self._esc(query), body_style))

        # Metrics
        metrics = result.get("metrics", [])
        if metrics:
            story.append(Paragraph("Key Metrics", heading_style))
            items = []
            for m in metrics:
                change = f" ({m.get('change_pct','')}% vs previous)" if m.get('change_pct') is not None else ""
                items.append(ListItem(Paragraph(f"<b>{self._esc(str(m.get('label','')))}</b>: {self._esc(str(m.get('value','')))}{change}", bullet_style)))
            story.append(ListFlowable(items, bulletType='bullet', start='•'))

        # Narrative
        narrative = result.get("narrative", "")
        if narrative:
            story.append(Paragraph("Analysis &amp; Recommendations", heading_style))
            self._render_narrative_pdf(story, narrative, body_style, heading_style, subhead_style, bullet_style)

        # Findings
        insights = result.get("insights", [])
        if insights:
            story.append(Paragraph("Key Findings", heading_style))
            for i in insights:
                sev = i.get("severity", "info").upper()
                conf = f" (Confidence: {int(i.get('confidence',0)*100)}%)" if i.get('confidence') else ""
                story.append(Paragraph(f"<b>[{sev}] {self._esc(i.get('title',''))}</b>{conf}", subhead_style))
                story.append(Paragraph(self._esc(i.get('body','')), bullet_style))
                if i.get('source'):
                    story.append(Paragraph(f"Method: {self._esc(i.get('source'))}", src_style))

        # Chart images
        if chart_images:
            self._add_charts_pdf(story, chart_images, heading_style, subhead_style, src_style)

        # Finance
        if finance:
            self._add_finance_pdf(story, finance, heading_style, body_style, bullet_style)

        # ── SIGNATURE BLOCK ──
        from reportlab.platypus import HRFlowable
        story.append(Spacer(1, 0.3*inch))
        story.append(HRFlowable(width="100%", thickness=0.7, color=HexColor('#cccccc'), spaceAfter=8))
        sig_label = ParagraphStyle('SigLabel', fontName='Times-Roman', fontSize=10,
                                   textColor=HexColor('#666666'), spaceAfter=2)
        sig_name = ParagraphStyle('SigName', fontName='Times-Bold', fontSize=12,
                                  textColor=HexColor('#0a4d4a'), spaceAfter=3)
        sig_line = ParagraphStyle('SigLine', fontName='Times-Roman', fontSize=10,
                                  textColor=HexColor('#333333'), spaceAfter=2)
        story.append(Paragraph("Prepared by", sig_label))
        story.append(Paragraph(self._esc(SIGNATURE["name"]), sig_name))
        story.append(Paragraph(f"Telephone: {SIGNATURE['phone1']}  |  {SIGNATURE['phone2']}", sig_line))
        story.append(Paragraph(f'LinkedIn: <link href="{SIGNATURE["linkedin"]}" color="#0a4d4a">{SIGNATURE["linkedin"]}</link>', sig_line))

        # (Platform footer intentionally omitted.)

        doc.build(story)
        buf.seek(0)
        return buf.read()

    def _add_pdf_table(self, story, header, body):
        """A real PDF table: header band, banded rows, numbers right-aligned."""
        from reportlab.platypus import Table, TableStyle, Spacer, Paragraph
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import inch
        if not header:
            return
        cell = ParagraphStyle('cell', fontName='Times-Roman', fontSize=8.5, leading=11)
        head = ParagraphStyle('head', fontName='Times-Bold', fontSize=8.5, leading=11,
                              textColor=colors.HexColor('#0a4d4a'))

        def is_num(v):
            try:
                float(str(v).replace(',', '').replace('%', ''))
                return True
            except Exception:
                return False

        data = [[Paragraph(str(h), head) for h in header]]
        for row in body:
            data.append([Paragraph(str(row[k]) if k < len(row) else '', cell)
                         for k in range(len(header))])

        # Share the page width evenly, but let a text column take more room
        avail = 6.6 * inch
        widths = [avail / len(header)] * len(header)

        t = Table(data, colWidths=widths, repeatRows=1)
        style = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eef6f5')),
            ('LINEBELOW', (0, 0), (-1, 0), 0.9, colors.HexColor('#0a4d4a')),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#d5dde3')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ]
        for r in range(1, len(data)):
            if r % 2 == 0:
                style.append(('BACKGROUND', (0, r), (-1, r), colors.HexColor('#f7fafb')))
        # Right-align columns whose body values are numeric
        for c in range(len(header)):
            vals = [row[c] for row in body if c < len(row)]
            if vals and all(is_num(v) or v in ('', '—') for v in vals):
                style.append(('ALIGN', (c, 1), (c, -1), 'RIGHT'))
        t.setStyle(TableStyle(style))
        story.append(t)
        story.append(Spacer(1, 10))

    def _render_narrative_pdf(self, story, narrative, body, heading, subhead, bullet):
        from reportlab.platypus import Paragraph, ListFlowable, ListItem
        import re
        narrative = re.sub(r'^\s*\{"narrative"\s*:\s*"', '', narrative)
        narrative = re.sub(r'",\s*"metrics"\s*:\s*\[.*$', '', narrative, flags=re.DOTALL)
        narrative = narrative.replace('\\n', '\n').replace('\\"', '"')

        bullet_buffer = []
        def flush_bullets():
            if bullet_buffer:
                story.append(ListFlowable([ListItem(Paragraph(b, bullet)) for b in bullet_buffer], bulletType='bullet', start='•'))
                bullet_buffer.clear()

        lines = narrative.split('\n')
        i = -1
        while i + 1 < len(lines):
            i += 1
            line = lines[i].strip()
            if not line:
                continue
            # ── Markdown table -> a real PDF table
            if (re.match(r'^\|.*\|$', line) and i + 1 < len(lines)
                    and re.match(r'^\|[\s:\-\|]+\|$', lines[i + 1].strip())):
                flush_bullets()
                header = [c.strip() for c in line.strip().strip('|').split('|')]
                j = i + 2
                body_rows = []
                while j < len(lines) and re.match(r'^\|.*\|$', lines[j].strip()):
                    body_rows.append([c.strip() for c in lines[j].strip().strip('|').split('|')])
                    j += 1
                self._add_pdf_table(story, header, body_rows)
                i = j - 1
                continue
            if line.startswith('## '):
                flush_bullets()
                story.append(Paragraph(self._esc(line[3:].strip()), heading))
            elif line.startswith('### '):
                flush_bullets()
                story.append(Paragraph(self._esc(line[4:].strip()), subhead))
            elif line.startswith(('- ', '• ', '* ')) or re.match(r'^\d+\.\s', line):
                clean = re.sub(r'^[-•*]\s*', '', line)
                clean = re.sub(r'^\d+\.\s*', '', clean)
                clean = self._esc(clean)
                clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean)
                bullet_buffer.append(clean)
            else:
                flush_bullets()
                text = self._esc(line)
                text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
                story.append(Paragraph(text, body))
        flush_bullets()

    def _add_charts_pdf(self, story, chart_images, heading, subhead, src):
        import base64, io
        from reportlab.platypus import Paragraph, Spacer, Image
        from reportlab.lib.units import inch
        story.append(Paragraph("Dashboard Visualisations", heading))
        for ci in chart_images[:8]:
            if not ci or not ci.get("image"):
                continue
            if ci.get("title"):
                story.append(Paragraph(self._esc(ci["title"]), subhead))
            try:
                img_bytes = base64.b64decode(ci["image"].split(",")[-1])
                img = Image(io.BytesIO(img_bytes), width=6*inch, height=3*inch, kind='proportional')
                story.append(img)
            except Exception:
                pass
            if ci.get("subtitle"):
                story.append(Paragraph(self._esc(ci["subtitle"]), src))
            story.append(Spacer(1, 0.15*inch))

    def _add_finance_pdf(self, story, finance, heading, body, bullet):
        from reportlab.platypus import Paragraph, ListFlowable, ListItem
        tax = finance.get("tax")
        acct = finance.get("accounting")
        fraud = finance.get("fraud")

        if tax and not tax.get("error"):
            story.append(Paragraph("Tax Analysis", heading))
            items = [ListItem(Paragraph(f"<b>{self._esc(str(m.get('label','')))}</b>: {self._esc(str(m.get('value','')))} (benchmark: {self._esc(str(m.get('benchmark','')))})", bullet)) for m in tax.get("metrics", [])]
            if items:
                story.append(ListFlowable(items, bulletType='bullet', start='•'))

        if acct and not acct.get("error"):
            story.append(Paragraph("Accounting Analysis", heading))
            story.append(Paragraph(f"Balance Sheet Health Score: {acct.get('health_score','?')}/100", body))
            items = [ListItem(Paragraph(f"<b>{self._esc(str(m.get('label','')))}</b>: {self._esc(str(m.get('value','')))}", bullet)) for m in acct.get("metrics", [])]
            if items:
                story.append(ListFlowable(items, bulletType='bullet', start='•'))

        if fraud and not fraud.get("error"):
            story.append(Paragraph("Fraud Detection", heading))
            story.append(Paragraph(f"Risk Score: {fraud.get('risk_score','?')}/100 ({fraud.get('risk_level','?')} risk)", body))

        actions = finance.get("priority_actions", [])
        if actions:
            story.append(Paragraph("Priority Actions", heading))
            items = [ListItem(Paragraph(f"<b>[{self._esc(str(a.get('module','')))}]</b> {self._esc(str(a.get('action','')))} — {self._esc(str(a.get('reason','')))}", bullet)) for a in actions]
            story.append(ListFlowable(items, bulletType='bullet', start='•'))

    def _esc(self, text):
        if not isinstance(text, str):
            text = str(text)
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


    # ── EXCEL EXPORT ──────────────────────────────────────────────────────────

    def build_excel_report(self, result: dict, finance: dict = None, chart_images: list = None,
                           doc_title: str = None, doc_subtitle: str = None,
                           line_spacing: float = 1.5) -> bytes:
        """
        Professional multi-sheet workbook:
          • Summary   — title block, question, the AI narrative (tables extracted)
          • Metrics   — KPI table
          • Insights  — findings table
          • Data      — the full dataset (falls back to raw_data_preview), styled
          • Finance   — tax / accounting / fraud / priority actions (when present)
          • Charts    — the dashboard chart images embedded
        Styled headers, frozen header rows, auto-sized columns, numbers as numbers.
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        import re, base64, io as _io

        # ── House style ──
        BRAND = "0A4D4A"
        BRAND_FILL = PatternFill("solid", fgColor="0A4D4A")
        HEAD_FILL = PatternFill("solid", fgColor="EEF6F5")
        BAND_FILL = PatternFill("solid", fgColor="F7FAFB")
        WHITE = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        HEAD_FONT = Font(name="Calibri", size=10, bold=True, color=BRAND)
        BODY_FONT = Font(name="Calibri", size=10, color="222222")
        TITLE_FONT = Font(name="Calibri", size=18, bold=True, color=BRAND)
        SUB_FONT = Font(name="Calibri", size=12, bold=True, color="222222")
        META_FONT = Font(name="Calibri", size=9, color="666666")
        ITAL = Font(name="Calibri", size=9, italic=True, color="666666")
        thin = Side(style="thin", color="D5DDE3")
        BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
        LEFT = Alignment(horizontal="left", vertical="center", wrap_text=True)
        RIGHT = Alignment(horizontal="right", vertical="center")
        CENTER = Alignment(horizontal="center", vertical="center")

        industry = str(result.get("industry", "General")).replace("_", " ").title()
        query = result.get("query", "Data Analysis")
        date_str = datetime.now(timezone.utc).strftime("%d %B %Y")

        wb = Workbook()

        def _num(v):
            """Return a float if the value is numeric, else None."""
            if isinstance(v, bool):
                return None
            if isinstance(v, (int, float)):
                return float(v)
            if v is None:
                return None
            s = str(v).strip().replace(",", "")
            pct = s.endswith("%")
            if pct:
                s = s[:-1]
            try:
                f = float(s)
                return f
            except Exception:
                return None

        def _autosize(ws, max_w=60, min_w=10):
            widths = {}
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    col = cell.column_letter
                    ln = max(len(x) for x in str(cell.value).split("\n"))
                    widths[col] = max(widths.get(col, 0), ln)
            for col, w in widths.items():
                ws.column_dimensions[col].width = max(min_w, min(max_w, w + 3))

        def _style_table(ws, header_row, first_data_row, last_row, ncols,
                         numeric_cols=None):
            numeric_cols = numeric_cols or set()
            # Header band
            for c in range(1, ncols + 1):
                cell = ws.cell(row=header_row, column=c)
                cell.font = HEAD_FONT
                cell.fill = HEAD_FILL
                cell.alignment = LEFT
                cell.border = BORDER
            # Body
            for r in range(first_data_row, last_row + 1):
                banded = (r - first_data_row) % 2 == 1
                for c in range(1, ncols + 1):
                    cell = ws.cell(row=r, column=c)
                    cell.font = BODY_FONT
                    cell.border = BORDER
                    if banded:
                        cell.fill = BAND_FILL
                    cell.alignment = RIGHT if c in numeric_cols else LEFT
            ws.freeze_panes = ws.cell(row=first_data_row, column=1)

        # ══ SHEET 1: SUMMARY ══════════════════════════════════════════════════
        ws = wb.active
        ws.title = "Summary"
        ws.sheet_view.showGridLines = False
        ws["A1"] = str(doc_title) if doc_title else f"{industry} Analysis"
        ws["A1"].font = TITLE_FONT
        ws["A2"] = str(doc_subtitle) if doc_subtitle else ""
        ws["A2"].font = SUB_FONT
        ws["A3"] = date_str
        ws["A3"].font = META_FONT
        ws["A5"] = "Analysis Question"
        ws["A5"].font = Font(name="Calibri", size=12, bold=True, color=BRAND)
        ws["A6"] = str(query)
        ws["A6"].font = BODY_FONT
        ws["A6"].alignment = LEFT

        r = 8
        narrative = result.get("narrative", "") or ""
        if narrative:
            ws.cell(row=r, column=1, value="Analysis & Recommendations").font = \
                Font(name="Calibri", size=12, bold=True, color=BRAND)
            r += 1
            # Clean JSON leakage the same way the Word/PDF builders do
            narrative = re.sub(r'^\s*\{"narrative"\s*:\s*"', '', narrative)
            narrative = re.sub(r'",\s*"metrics"\s*:\s*\[.*$', '', narrative, flags=re.DOTALL)
            narrative = narrative.replace('\\n', '\n').replace('\\"', '"')
            for raw in narrative.split("\n"):
                line = raw.rstrip()
                if not line.strip():
                    continue
                # Skip markdown table rows in the summary prose — they'd read as pipes.
                if re.match(r'^\|.*\|$', line.strip()):
                    continue
                txt = line
                bold = False
                if txt.strip().startswith("## "):
                    txt = txt.strip()[3:]; bold = True
                elif txt.strip().startswith("### "):
                    txt = txt.strip()[4:]; bold = True
                txt = re.sub(r'\*\*(.+?)\*\*', r'\1', txt)
                cell = ws.cell(row=r, column=1, value=txt)
                cell.font = Font(name="Calibri", size=11, bold=True, color=BRAND) if bold \
                    else BODY_FONT
                cell.alignment = LEFT
                r += 1
        ws.column_dimensions["A"].width = 110

        # ══ SHEET 2: METRICS ══════════════════════════════════════════════════
        metrics = result.get("metrics", []) or []
        if metrics:
            ms = wb.create_sheet("Metrics")
            ms.sheet_view.showGridLines = False
            hdr = ["Metric", "Value", "Change vs previous", "Benchmark"]
            for c, h in enumerate(hdr, 1):
                ms.cell(row=1, column=c, value=h)
            for i, m in enumerate(metrics, start=2):
                ms.cell(row=i, column=1, value=str(m.get("label", "")))
                v = _num(m.get("value"))
                ms.cell(row=i, column=2, value=v if v is not None else str(m.get("value", "")))
                cp = m.get("change_pct")
                ms.cell(row=i, column=3,
                        value=(f"{cp}%" if cp is not None else ""))
                ms.cell(row=i, column=4, value=str(m.get("benchmark", "") or ""))
            _style_table(ms, 1, 2, len(metrics) + 1, 4, numeric_cols={2})
            _autosize(ms)

        # ══ SHEET 3: INSIGHTS ═════════════════════════════════════════════════
        insights = result.get("insights", []) or []
        if insights:
            ins = wb.create_sheet("Insights")
            ins.sheet_view.showGridLines = False
            hdr = ["Severity", "Finding", "Detail", "Confidence", "Method"]
            for c, h in enumerate(hdr, 1):
                ins.cell(row=1, column=c, value=h)
            for i, it in enumerate(insights, start=2):
                ins.cell(row=i, column=1, value=str(it.get("severity", "info")).upper())
                ins.cell(row=i, column=2, value=str(it.get("title", "")))
                ins.cell(row=i, column=3, value=str(it.get("body", "")))
                conf = it.get("confidence")
                ins.cell(row=i, column=4,
                         value=(f"{int(conf*100)}%" if conf else ""))
                ins.cell(row=i, column=5, value=str(it.get("source", "") or ""))
            _style_table(ins, 1, 2, len(insights) + 1, 5)
            # Detail column is long — give it room and wrap
            ins.column_dimensions["C"].width = 70
            ins.column_dimensions["B"].width = 34
            ins.column_dimensions["E"].width = 30
            _autosize(ins, max_w=70)
            ins.column_dimensions["C"].width = 70  # keep after autosize

        # ══ SHEET 4: DATA ═════════════════════════════════════════════════════
        # Prefer the full dataset the frontend may send; fall back to the preview.
        data_map = None
        data_rows = result.get("data")
        if not isinstance(data_rows, list) or not data_rows:
            data_rows = result.get("raw_data_preview") or []
        if isinstance(data_rows, list) and data_rows and isinstance(data_rows[0], dict):
            ds = wb.create_sheet("Data")
            ds.sheet_view.showGridLines = False
            # Union of keys, preserving first-seen order
            cols = []
            for row in data_rows[:5000]:
                for k in row.keys():
                    if k not in cols:
                        cols.append(k)
            for c, k in enumerate(cols, 1):
                ds.cell(row=1, column=c, value=str(k).replace("_", " "))
            numeric_cols = set()
            MAX_ROWS = 150000
            for i, row in enumerate(data_rows[:MAX_ROWS], start=2):
                for c, k in enumerate(cols, 1):
                    v = row.get(k)
                    nv = _num(v)
                    if nv is not None and not (isinstance(v, str) and v.strip() == ""):
                        ds.cell(row=i, column=c, value=nv)
                        numeric_cols.add(c)
                    else:
                        ds.cell(row=i, column=c, value=("" if v is None else str(v)))
            last = min(len(data_rows), MAX_ROWS) + 1
            # A column is "numeric" only if it was numeric in every populated cell;
            # recompute cleanly to avoid right-aligning mixed columns.
            clean_numeric = set()
            for c in range(1, len(cols) + 1):
                saw_val = False
                all_num = True
                for r in range(2, last + 1):
                    val = ds.cell(row=r, column=c).value
                    if val == "" or val is None:
                        continue
                    saw_val = True
                    if not isinstance(val, (int, float)):
                        all_num = False
                        break
                if saw_val and all_num:
                    clean_numeric.add(c)
            # Record where each field lives so the Dashboard can write live
            # formulas that point back at this sheet.
            data_map = {"cols": list(cols), "first_row": 2, "last_row": last,
                        "total_rows": len(data_rows), "written_rows": min(len(data_rows), MAX_ROWS),
                        "truncated": len(data_rows) > MAX_ROWS}
            # Style the HEADER only. Styling every data cell (border+font+fill)
            # is what makes large workbooks slow to write and heavy to open, so
            # for the data body we write values only and rely on the frozen
            # header + column widths for readability.
            for c in range(1, len(cols) + 1):
                hc = ds.cell(row=1, column=c)
                hc.font = HEAD_FONT
                hc.fill = HEAD_FILL
                hc.alignment = LEFT
                hc.border = BORDER
            ds.freeze_panes = "A2"
            ds.auto_filter.ref = f"A1:{get_column_letter(len(cols))}{last}"
            for c in clean_numeric:
                ds.cell(row=2, column=c).number_format = "#,##0.##"
            _autosize(ds, max_w=40)
            if len(data_rows) > MAX_ROWS:
                note = ds.cell(row=last + 2, column=1,
                               value=f"Showing the first {MAX_ROWS:,} of {len(data_rows):,} rows.")
                note.font = ITAL

        # ══ ANALYTICAL DASHBOARD (first sheet) ══════════════════════════════
        dash_info = None
        try:
            dash_info = self._build_dashboard_sheet(wb, result,
                                                    data_rows if isinstance(data_rows, list) else [],
                                                    data_map=data_map)
        except Exception as e:
            logger.warning(f"Dashboard sheet skipped: {e}")
        try:
            self._add_methodology_sheet(wb, result, data_map,
                                        (dash_info or {}).get("value_col"),
                                        (dash_info or {}).get("breakdown_cols"))
        except Exception as e:
            logger.warning(f"Methodology sheet skipped: {e}")

        # ══ SHEET 5: FINANCE ══════════════════════════════════════════════════
        if finance:
            self._add_finance_excel(wb, finance, HEAD_FONT, HEAD_FILL, BODY_FONT,
                                    BRAND, BORDER, LEFT, RIGHT, _num, _autosize, _style_table)

        # ══ SHEET 6: CHARTS ═══════════════════════════════════════════════════
        if chart_images:
            try:
                from openpyxl.drawing.image import Image as XLImage
                cs = wb.create_sheet("Charts")
                cs.sheet_view.showGridLines = False
                cs.column_dimensions["A"].width = 14
                anchor_row = 1
                placed = 0
                for ci in chart_images[:12]:
                    if not ci or not ci.get("image"):
                        continue
                    title = ci.get("title", "") or ""
                    if title:
                        tcell = cs.cell(row=anchor_row, column=1, value=title)
                        tcell.font = Font(name="Calibri", size=12, bold=True, color=BRAND)
                        anchor_row += 1
                    if ci.get("subtitle"):
                        scell = cs.cell(row=anchor_row, column=1, value=ci["subtitle"])
                        scell.font = ITAL
                        anchor_row += 1
                    try:
                        b64 = ci["image"].split(",")[-1]
                        img_bytes = base64.b64decode(b64)
                        bio = _io.BytesIO(img_bytes)
                        img = XLImage(bio)
                        # Scale to a sensible width (~640px) keeping aspect ratio
                        if img.width and img.height:
                            target_w = 640
                            ratio = target_w / float(img.width)
                            img.width = target_w
                            img.height = int(img.height * ratio)
                        img.anchor = f"A{anchor_row}"
                        cs.add_image(img)
                        # ~20px per row → advance enough rows to clear the image
                        rows_tall = int((img.height or 320) / 18) + 3
                        anchor_row += rows_tall
                        placed += 1
                    except Exception as e:
                        logger.warning(f"Excel chart embed failed: {e}")
                        anchor_row += 2
                if not placed:
                    # No image decoded — drop the empty sheet to avoid confusion
                    wb.remove(cs)
            except Exception as e:
                logger.warning(f"Excel charts sheet skipped: {e}")

        # Make sure Summary is the active/first sheet on open
        wb.active = 0

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf.read()

    # ── ANALYTICAL DASHBOARD SHEET ────────────────────────────────────────────
    # Builds an executive summary sheet like a hand-made analyst workbook:
    # a metrics block, several "Top-N by value" breakdown tables, and a recency
    # cross-tab with grouped headers — all computed from the actual data.

    def _build_dashboard_sheet(self, wb, result: dict, data_rows: list, data_map: dict = None):
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        try:
            import pandas as pd
        except Exception:
            return  # pandas needed for the aggregations; skip cleanly if absent
        if not data_rows or not isinstance(data_rows[0], dict):
            return

        df = pd.DataFrame(data_rows)
        if df.empty:
            return

        # ── Palette (executive blue, matching a classic analyst workbook) ──
        BAND = PatternFill("solid", fgColor="4472C4")     # header band
        BAND2 = PatternFill("solid", fgColor="8FAADC")    # sub-header / group band
        BAND_GREY = PatternFill("solid", fgColor="A6A6A6") # secondary group band
        LIGHT = PatternFill("solid", fgColor="D9E1F2")    # metric label fill
        TOTAL_FILL = PatternFill("solid", fgColor="DDEBF7")
        WHITE_B = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
        HEAD_B = Font(name="Calibri", size=10, bold=True, color="1F3864")
        BODY = Font(name="Calibri", size=10, color="000000")
        BODY_B = Font(name="Calibri", size=10, bold=True, color="000000")
        TITLE = Font(name="Calibri", size=15, bold=True, color="1F3864")
        SUBT = Font(name="Calibri", size=9, italic=True, color="595959")
        thin = Side(style="thin", color="BFBFBF")
        BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
        L = Alignment(horizontal="left", vertical="center")
        R = Alignment(horizontal="right", vertical="center")
        C = Alignment(horizontal="center", vertical="center")

        ws = wb.create_sheet("Dashboard")
        ws.sheet_view.showGridLines = False

        # ── Classify columns ──
        num_cols, cat_cols = [], []
        for c in df.columns:
            s = pd.to_numeric(df[c], errors="coerce")
            if s.notna().sum() >= max(3, 0.5 * df[c].notna().sum()):
                num_cols.append(c)
            else:
                nun = df[c].nunique(dropna=True)
                if 0 < nun <= max(50, len(df) // 2):
                    cat_cols.append(c)

        def numify(c):
            return pd.to_numeric(df[c], errors="coerce")

        # The primary value column: prefer money-ish/aggregate names, else biggest sum
        def pick_value_col():
            if not num_cols:
                return None
            pref = [c for c in num_cols if any(w in str(c).lower()
                    for w in ("value", "sales", "revenue", "amount", "total", "gross",
                              "net", "price", "cost", "spend", "gmv"))]
            skip = ("rank", "id", "index", "year", "no", "number", "code", "qty", "count")
            base = pref or [c for c in num_cols if not any(k in str(c).lower() for k in skip)] or num_cols
            sums = {c: float(numify(c).sum()) for c in base}
            return max(sums, key=sums.get) if sums else num_cols[0]

        value_col = pick_value_col()
        # A count-ish column (units/volume/qty) for a secondary measure, if present
        count_col = next((c for c in num_cols if any(w in str(c).lower()
                          for w in ("volume", "units", "qty", "quantity", "count"))), None)

        industry = str(result.get("industry", "General")).replace("_", " ").title()
        date_str = datetime.now(timezone.utc).strftime("%d %B %Y")

        def fmt(v):
            try:
                return f"{float(v):,.0f}"
            except Exception:
                return str(v)

        # ── Title band ──
        ws.merge_cells("A1:F1")
        ws["A1"] = f"{industry} Analysis — Overview"
        ws["A1"].font = TITLE
        ws.merge_cells("A2:F2")
        cur = " · Values summarised" + (f" by {value_col}" if value_col else "")
        ws["A2"] = f"{len(df):,} records · {date_str}{cur}"
        ws["A2"].font = SUBT
        row = 4

        # ── Helper: write a titled table with a coloured header band ──
        def section_header(r, c0, span, text, fill=BAND, font=WHITE_B):
            for i in range(span):
                cell = ws.cell(row=r, column=c0 + i)
                cell.fill = fill
                cell.border = BORDER
                if i == 0:
                    cell.value = text
                    cell.font = font
                    cell.alignment = L
            return r + 1

        def col_headers(r, c0, headers):
            for i, h in enumerate(headers):
                cell = ws.cell(row=r, column=c0 + i, value=h)
                cell.font = HEAD_B
                cell.fill = LIGHT
                cell.border = BORDER
                cell.alignment = L if i == 0 else R
            return r + 1

        def data_row(r, c0, cells, numeric_from=1, bold=False, fill=None):
            for i, v in enumerate(cells):
                cell = ws.cell(row=r, column=c0 + i)
                is_num = isinstance(v, (int, float)) and not isinstance(v, bool)
                cell.value = v
                cell.font = BODY_B if bold else BODY
                cell.border = BORDER
                cell.alignment = R if i >= numeric_from else L
                if is_num:
                    cell.number_format = "#,##0"
                if fill:
                    cell.fill = fill
            return r + 1

        # ── KEY METRICS block (left) ──
        metrics = []
        metrics.append((f"Total Records", len(df)))
        if value_col is not None:
            metrics.append((f"Total {value_col.replace('_',' ').title()}", float(numify(value_col).sum())))
        if count_col is not None:
            metrics.append((f"Total {count_col.replace('_',' ').title()}", float(numify(count_col).sum())))
        # a rate-ish column → show its mean
        rate_col = next((c for c in num_cols if any(w in str(c).lower()
                        for w in ("margin", "rate", "pct", "percent", "ratio"))), None)
        if rate_col is not None:
            metrics.append((f"Average {rate_col.replace('_',' ').title()}",
                            round(float(numify(rate_col).mean()), 1)))
        for c in cat_cols[:3]:
            metrics.append((f"Distinct {c.replace('_',' ').title()}", int(df[c].nunique(dropna=True))))

        km_row = section_header(row, 1, 2, "KEY METRICS")
        for label, val in metrics:
            cell = ws.cell(row=km_row, column=1, value=label)
            cell.font = BODY; cell.fill = LIGHT; cell.border = BORDER; cell.alignment = L
            vcell = ws.cell(row=km_row, column=2, value=val)
            vcell.font = BODY_B; vcell.border = BORDER; vcell.alignment = R
            if isinstance(val, (int, float)) and not isinstance(val, bool):
                vcell.number_format = "#,##0.#" if isinstance(val, float) else "#,##0"
            km_row += 1
        metrics_bottom = km_row

        # ── Breakdown tables: Top-N by value for the best categorical columns ──
        # Rank categoricals by how concentrated their value is (more explanatory).
        breakdown_cols = []
        if value_col is not None:
            scored = []
            for c in cat_cols:
                try:
                    g = df.groupby(c)[value_col].apply(lambda s: pd.to_numeric(s, errors="coerce").sum())
                    if g.sum() > 0 and 1 < len(g) <= 60:
                        scored.append((c, len(g)))
                except Exception:
                    continue
            # Prefer dimension-like names first, then by cardinality
            def rank_key(item):
                c, n = item
                pri = 0 if any(w in str(c).lower() for w in
                       ("brand", "category", "style", "source", "product", "region",
                        "segment", "type", "class", "group")) else 1
                return (pri, -n if n <= 20 else n)
            scored.sort(key=rank_key)
            breakdown_cols = [c for c, _ in scored[:4]]

        # Lay tables starting under the metrics block, left column c0=1
        tbl_row = metrics_bottom + 2
        total_value = float(numify(value_col).sum()) if value_col is not None else 0

        for c in breakdown_cols:
            g = df.groupby(c).agg(
                _cnt=(value_col, "size"),
                _val=(value_col, lambda s: pd.to_numeric(s, errors="coerce").sum()),
            ).reset_index().rename(columns={c: "name"})
            g = g.sort_values("_val", ascending=False).head(15)
            title = f"By {c.replace('_',' ').title()}"
            # Can we point live formulas at the Data sheet for this pairing?
            live = None
            if data_map and c in (data_map.get("cols") or []) and value_col in (data_map.get("cols") or []):
                k_col = get_column_letter(data_map["cols"].index(c) + 1)
                v_col = get_column_letter(data_map["cols"].index(value_col) + 1)
                live = {"k": k_col, "v": v_col,
                        "r1": data_map["first_row"], "r2": data_map["last_row"]}
            span = 5 if live else 4
            tbl_row = section_header(tbl_row, 1, span, title.upper())
            heads = [c.replace('_', ' ').title(), "Count", f"{value_col.replace('_',' ').title()}", "% of Value"]
            if live:
                heads.append("Formula used")
            tbl_row = col_headers(tbl_row, 1, heads)
            shown_val = 0
            for _, rrow in g.iterrows():
                pct = (rrow["_val"] / total_value * 100) if total_value else 0
                shown_val += rrow["_val"]
                r = tbl_row
                name = str(rrow["name"])
                if live:
                    # LIVE formulas: these recalculate if the Data sheet is edited.
                    rng_k = f"Data!${live['k']}${live['r1']}:${live['k']}${live['r2']}"
                    rng_v = f"Data!${live['v']}${live['r1']}:${live['v']}${live['r2']}"
                    safe = name.replace('"', '""')
                    cnt_f = f'=COUNTIF({rng_k},"{safe}")'
                    sum_f = f'=SUMIF({rng_k},"{safe}",{rng_v})'
                    ws.cell(row=r, column=1, value=name).font = BODY
                    ws.cell(row=r, column=1).border = BORDER
                    ws.cell(row=r, column=1).alignment = L
                    cc = ws.cell(row=r, column=2, value=cnt_f)
                    cc.font = BODY; cc.border = BORDER; cc.alignment = R; cc.number_format = "#,##0"
                    vc = ws.cell(row=r, column=3, value=sum_f)
                    vc.font = BODY; vc.border = BORDER; vc.alignment = R; vc.number_format = "#,##0"
                    # Show the formula as text so the method is visible on the page
                    fc = ws.cell(row=r, column=5, value=sum_f.replace("=", "", 1))
                    fc.font = SUBT; fc.border = BORDER; fc.alignment = L
                else:
                    data_row(r, 1, [name, int(rrow["_cnt"]), round(float(rrow["_val"]), 0)])
                pcell = ws.cell(row=r, column=4,
                                value=(f"=IFERROR(C{r}/SUM($C${tbl_row - len(g)}:$C${tbl_row + len(g)}),0)" if False else pct / 100))
                pcell.number_format = "0.0%"; pcell.font = BODY; pcell.border = BORDER; pcell.alignment = R
                tbl_row += 1
            # Totals row
            tot_pct = (shown_val / total_value) if total_value else 0
            data_row(tbl_row, 1, ["Total", int(g["_cnt"].sum()), round(float(shown_val), 0)],
                     bold=True, fill=TOTAL_FILL)
            tcell = ws.cell(row=tbl_row, column=4, value=tot_pct)
            tcell.number_format = "0.0%"; tcell.font = BODY_B; tcell.border = BORDER
            tcell.alignment = R; tcell.fill = TOTAL_FILL
            tbl_row += 2

        # ── Recency / split cross-tab (grouped headers) when a split exists ──
        # Find a low-cardinality column that splits into 2–3 buckets (or derive
        # one from a days-since column), then show value/count per bucket per
        # top category — the "Sold within 30 / Sold > 30" style block.
        split_col = None
        split_buckets = None
        days_col = next((c for c in num_cols if any(w in str(c).lower()
                        for w in ("days", "age", "recency", "last_sold", "dayssince"))), None)
        if days_col is not None:
            d = numify(days_col)
            bucket = pd.Series(pd.cut(d, bins=[-1, 30, 10**9], labels=["Within 30 days", "Over 30 days"]))
            if bucket.notna().sum() > 0:
                split_col = "_recency_"
                df = df.assign(**{split_col: bucket.values})
                split_buckets = ["Within 30 days", "Over 30 days"]
        if split_col is None:
            for c in cat_cols:
                nun = df[c].nunique(dropna=True)
                if 2 <= nun <= 3:
                    split_col = c
                    split_buckets = [str(x) for x in df[c].dropna().unique()][:3]
                    break

        if split_col is not None and value_col is not None and breakdown_cols:
            dim = breakdown_cols[0]
            tbl_row = section_header(tbl_row, 1, 1 + 3 * len(split_buckets),
                                     f"{dim.replace('_',' ').upper()} × {('RECENCY' if split_col=='_recency_' else split_col.replace('_',' ').upper())}")
            # Grouped header band: dim | [bucket1 spanning 3] | [bucket2 spanning 3]
            gh = tbl_row
            ws.cell(row=gh, column=1, value=dim.replace('_', ' ').title()).font = HEAD_B
            ws.cell(row=gh, column=1).fill = LIGHT
            ws.cell(row=gh, column=1).border = BORDER
            fills = [BAND2, BAND_GREY, BAND2]
            for bi, b in enumerate(split_buckets):
                c0 = 2 + bi * 3
                ws.merge_cells(start_row=gh, start_column=c0, end_row=gh, end_column=c0 + 2)
                gcell = ws.cell(row=gh, column=c0, value=str(b))
                gcell.font = WHITE_B; gcell.fill = fills[bi % len(fills)]
                gcell.alignment = C; gcell.border = BORDER
                for k in range(1, 3):
                    ws.cell(row=gh, column=c0 + k).fill = fills[bi % len(fills)]
                    ws.cell(row=gh, column=c0 + k).border = BORDER
            # Sub-header row
            sh = gh + 1
            ws.cell(row=sh, column=1, value="").border = BORDER
            for bi in range(len(split_buckets)):
                c0 = 2 + bi * 3
                for k, h in enumerate([f"{value_col.replace('_',' ').title()}", "Count", (count_col.replace('_',' ').title() if count_col else "Volume")]):
                    cc = ws.cell(row=sh, column=c0 + k, value=h)
                    cc.font = HEAD_B; cc.fill = LIGHT; cc.border = BORDER; cc.alignment = R
            # Body: top 12 of the dimension
            top_dim = (df.groupby(dim)[value_col]
                         .apply(lambda s: pd.to_numeric(s, errors="coerce").sum())
                         .sort_values(ascending=False).head(12).index.tolist())
            br = sh + 1
            for name in top_dim:
                sub = df[df[dim] == name]
                ws.cell(row=br, column=1, value=str(name)).font = BODY
                ws.cell(row=br, column=1).border = BORDER
                ws.cell(row=br, column=1).alignment = L
                for bi, b in enumerate(split_buckets):
                    seg = sub[sub[split_col].astype(str) == str(b)]
                    val = float(pd.to_numeric(seg[value_col], errors="coerce").sum())
                    cnt = int(len(seg))
                    vol = float(pd.to_numeric(seg[count_col], errors="coerce").sum()) if count_col else 0
                    c0 = 2 + bi * 3
                    for k, v in enumerate([round(val, 0), cnt, round(vol, 0)]):
                        cc = ws.cell(row=br, column=c0 + k, value=v)
                        cc.font = BODY; cc.border = BORDER; cc.alignment = R
                        cc.number_format = "#,##0"
                br += 1
            tbl_row = br + 1

        # ── Column widths ──
        ws.column_dimensions["A"].width = 30
        for col in ["B", "C", "D", "E", "F", "G", "H", "I", "J"]:
            ws.column_dimensions[col].width = 15

        # Make Dashboard the first sheet
        wb.move_sheet("Dashboard", -(len(wb.sheetnames) - 1))
        return {"value_col": value_col, "breakdown_cols": breakdown_cols}

    def _add_methodology_sheet(self, wb, result: dict, data_map: dict = None,
                               value_col=None, breakdown_cols=None):
        """
        Documents every calculation in the workbook: the live Excel formulas,
        and the statistical methods (computed in Python, not reproducible as
        cell formulas). Also records any row truncation so the figures are
        never silently misleading.
        """
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        BRAND = "1F3864"
        HEAD_FILL = PatternFill("solid", fgColor="D9E1F2")
        HEAD = Font(name="Calibri", size=10, bold=True, color=BRAND)
        BODY = Font(name="Calibri", size=10)
        MONO = Font(name="Consolas", size=9.5)
        TITLE = Font(name="Calibri", size=14, bold=True, color=BRAND)
        WARN = Font(name="Calibri", size=10, bold=True, color="C00000")
        thin = Side(style="thin", color="BFBFBF")
        BORDER = Border(left=thin, right=thin, top=thin, bottom=thin)
        L = Alignment(horizontal="left", vertical="top", wrap_text=True)

        ms = wb.create_sheet("Methodology")
        ms.sheet_view.showGridLines = False
        ms["A1"] = "Methodology & Formulas"
        ms["A1"].font = TITLE
        ms["A2"] = "How every figure in this workbook was calculated."
        ms["A2"].font = Font(name="Calibri", size=9, italic=True, color="595959")
        r = 4

        # Truncation warning — critical for honesty about live formulas
        if data_map and data_map.get("truncated"):
            ms.cell(row=r, column=1,
                    value=(f"NOTE: the Data sheet contains {data_map['written_rows']:,} of "
                           f"{data_map['total_rows']:,} total rows. Live Excel formulas on the "
                           f"Dashboard sum only the rows present in this workbook, so they may "
                           f"differ from figures computed over the full dataset.")).font = WARN
            ms.cell(row=r, column=1).alignment = L
            ms.merge_cells(start_row=r, start_column=1, end_row=r, end_column=3)
            ms.row_dimensions[r].height = 30
            r += 2

        # Table header
        for c, h in enumerate(["What", "How it was calculated", "Excel formula (if live)"], 1):
            cell = ms.cell(row=r, column=c, value=h)
            cell.font = HEAD; cell.fill = HEAD_FILL; cell.border = BORDER; cell.alignment = L
        r += 1

        def row(what, how, formula=""):
            nonlocal r
            ms.cell(row=r, column=1, value=what).font = BODY
            ms.cell(row=r, column=2, value=how).font = BODY
            fc = ms.cell(row=r, column=3, value=formula)
            fc.font = MONO if formula else BODY
            for c in range(1, 4):
                ms.cell(row=r, column=c).border = BORDER
                ms.cell(row=r, column=c).alignment = L
            r += 1

        cols = (data_map or {}).get("cols") or []
        vcol = value_col or "the value column"

        # Dashboard calculations
        if breakdown_cols:
            for c in breakdown_cols:
                nice = str(c).replace("_", " ").title()
                if c in cols and value_col in cols:
                    from openpyxl.utils import get_column_letter as _gl
                    k = _gl(cols.index(c) + 1); v = _gl(cols.index(value_col) + 1)
                    r1, r2 = data_map["first_row"], data_map["last_row"]
                    row(f"By {nice} — Count",
                        f"Number of records in each {nice} group.",
                        f'COUNTIF(Data!${k}${r1}:${k}${r2}, "<group>")')
                    row(f"By {nice} — {str(vcol).replace('_',' ').title()}",
                        f"Sum of {vcol} for each {nice} group. Live formula — recalculates if the Data sheet changes.",
                        f'SUMIF(Data!${k}${r1}:${k}${r2}, "<group>", Data!${v}${r1}:${v}${r2})')
                else:
                    row(f"By {nice}", f"Sum of {vcol} grouped by {nice}, computed in Python (pandas groupby).", "")
            row("% of Value", "Each group's value divided by the total across all groups shown.", "group value / total")

        row("Key Metrics — totals",
            f"Column sums and averages over the analysed rows (pandas).", "")
        row("Distinct counts", "Number of unique values in the column (pandas nunique).", "")
        row("Cross-tab (x Recency)",
            "Rows split into buckets by the days-since column (0-30, over 30), then value and count summed per bucket per group.", "")

        # Statistical methods — cannot be Excel formulas
        stats_used = set()
        for i in (result.get("insights") or []):
            src = str(i.get("source", "") or "")
            if src:
                stats_used.add(src.split("\u00b7")[0].strip())
        for src in sorted(x for x in stats_used if x):
            if "Z-score" in src:
                row("Anomaly detection", "Z-score: (value - mean) / standard deviation. Points beyond 3 sigma flagged as anomalies. Computed in Python (NumPy).", "")
            elif "Pearson" in src or "correlation" in src.lower():
                row("Correlations", "Pearson correlation coefficient with Bonferroni correction for multiple comparisons (SciPy).", "")
            elif "regression" in src.lower() or "trend" in src.lower():
                row("Trends", "Ordinary least squares linear regression; R-squared and p-value reported (SciPy/statsmodels).", "")
            elif src:
                row("Finding method", src, "")

        ms.column_dimensions["A"].width = 34
        ms.column_dimensions["B"].width = 68
        ms.column_dimensions["C"].width = 52
        return ms

    def _add_finance_excel(self, wb, finance, HEAD_FONT, HEAD_FILL, BODY_FONT,
                           BRAND, BORDER, LEFT, RIGHT, _num, _autosize, _style_table):
        from openpyxl.styles import Font, Alignment
        fs = wb.create_sheet("Finance")
        fs.sheet_view.showGridLines = False
        title_font = Font(name="Calibri", size=12, bold=True, color=BRAND)
        r = 1

        def section(title):
            nonlocal r
            cell = fs.cell(row=r, column=1, value=title)
            cell.font = title_font
            r += 1

        def table(header, rows, numeric_cols=None):
            nonlocal r
            if not rows:
                return
            header_row = r
            for c, h in enumerate(header, 1):
                fs.cell(row=header_row, column=c, value=h)
            first = r + 1
            for row in rows:
                for c, v in enumerate(row, 1):
                    nv = _num(v)
                    fs.cell(row=first, column=c,
                            value=nv if (nv is not None and str(v).strip() != "") else str(v))
                first += 1
            last = first - 1
            _style_table(fs, header_row, header_row + 1, last, len(header),
                         numeric_cols=numeric_cols or set())
            r = last + 2

        tax = finance.get("tax")
        acct = finance.get("accounting")
        fraud = finance.get("fraud")

        if tax and not tax.get("error"):
            section("Tax Analysis")
            table(["Metric", "Value", "Benchmark"],
                  [[m.get("label", ""), m.get("value", ""), m.get("benchmark", "")]
                   for m in tax.get("metrics", [])])
            findings = tax.get("findings", [])[:8]
            if findings:
                table(["Finding", "Detail"],
                      [[f.get("title", ""), f.get("body", "")] for f in findings])

        if acct and not acct.get("error"):
            section(f"Accounting Analysis — Health Score: {acct.get('health_score','?')}/100")
            table(["Metric", "Value"],
                  [[m.get("label", ""), m.get("value", "")] for m in acct.get("metrics", [])])

        if fraud and not fraud.get("error"):
            section(f"Fraud Detection — Risk: {fraud.get('risk_level','?')} ({fraud.get('risk_score','?')}/100)")
            findings = [f for f in fraud.get("findings", []) if f.get("severity") in ("critical", "warning")][:8]
            if findings:
                table(["Severity", "Finding", "Detail"],
                      [[f.get("severity", "").upper(), f.get("title", ""), f.get("body", "")]
                       for f in findings])

        actions = finance.get("priority_actions", [])
        if actions:
            section("Priority Actions")
            table(["Module", "Action", "Reason"],
                  [[a.get("module", ""), a.get("action", ""), a.get("reason", "")] for a in actions])

        _autosize(fs, max_w=70)


document_service = DocumentExportService()
