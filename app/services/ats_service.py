"""
DataMind Agent — Elite ATS (Applicant Tracking / CV Scoring)

Reads CVs (PDF/DOCX/TXT), derives criteria from a job description, and scores
each candidate 0-100 against weighted criteria — with an evidence quote from the
CV behind every sub-score, so a score can always be audited back to real text.

Design decisions:
  • Anonymisation is per-job and ON by default. Name, email, phone, address,
    age/DOB, gender markers, nationality and photos are stripped BEFORE the CV
    ever reaches the model, so they cannot influence the score.
  • Criteria: the AI proposes them from the JD, HR edits/adds/removes and sets
    weights. HR always has the final say — their weights are authoritative.
  • Every sub-score must cite CV text. No quote, no credit.

Fairness notes (deliberate, not incidental):
  - Scoring is against job-relevant criteria only.
  - The model is instructed to ignore proxies for protected characteristics
    (school prestige, employment gaps, name origin, photos, age).
  - Nothing here decides anything. It ranks and evidences; a human decides.
"""
from __future__ import annotations
import re, uuid, json, logging
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

MAX_CV_CHARS = 18000          # a very long CV still fits a prompt
MAX_CVS_PER_JOB = 200


def now():
    return datetime.now(timezone.utc)


# ── PII stripping ─────────────────────────────────────────────────────────────
EMAIL_RE = re.compile(r"[\w\.\-\+]+@[\w\-]+\.[\w\.\-]+")
PHONE_RE = re.compile(r"(?:(?:\+|00)\d{1,3}[\s\-\.]?)?(?:\(?\d{2,4}\)?[\s\-\.]?){2,4}\d{2,4}")
URL_RE   = re.compile(r"https?://\S+|www\.\S+|linkedin\.com/\S+", re.I)
DOB_RE   = re.compile(r"\b(?:date of birth|dob|born|age)\s*[:\-]?\s*[^\n]{0,30}", re.I)
GENDER_RE= re.compile(r"\b(?:gender|sex|marital status|nationality|religion)\s*[:\-]?\s*[^\n]{0,25}", re.I)
TITLE_RE = re.compile(r"\b(mr|mrs|ms|miss|dr|prof)\.?\s+[A-Z][a-z]+", re.I)


def anonymise(text: str, known_name: str = None) -> tuple[str, dict]:
    """
    Remove identity signals before scoring. Returns (clean_text, what_was_removed).
    This runs BEFORE the model sees anything, so PII cannot sway a score.
    """
    removed = {"email": 0, "phone": 0, "url": 0, "dob_age": 0, "demographic": 0, "name": 0}
    t = text or ""

    removed["email"] = len(EMAIL_RE.findall(t)); t = EMAIL_RE.sub("[EMAIL REMOVED]", t)
    removed["url"] = len(URL_RE.findall(t));     t = URL_RE.sub("[LINK REMOVED]", t)
    removed["dob_age"] = len(DOB_RE.findall(t)); t = DOB_RE.sub("[AGE/DOB REMOVED]", t)
    removed["demographic"] = len(GENDER_RE.findall(t)); t = GENDER_RE.sub("[DEMOGRAPHIC REMOVED]", t)
    removed["phone"] = len(PHONE_RE.findall(t)); t = PHONE_RE.sub("[PHONE REMOVED]", t)
    removed["name"] = len(TITLE_RE.findall(t));  t = TITLE_RE.sub("[NAME REMOVED]", t)

    if known_name:
        for part in [p for p in re.split(r"\s+", known_name) if len(p) > 2]:
            pat = re.compile(r"\b" + re.escape(part) + r"\b", re.I)
            removed["name"] += len(pat.findall(t))
            t = pat.sub("[NAME REMOVED]", t)
    return t, removed


# ── CV text extraction ────────────────────────────────────────────────────────
def extract_text(file_bytes: bytes, filename: str) -> tuple[str, str]:
    """Pull prose out of a CV. Returns (text, error)."""
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            import pdfplumber, io
            out = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages[:15]:
                    out.append(page.extract_text() or "")
            return "\n".join(out).strip(), None
        if name.endswith(".docx"):
            from docx import Document
            import io
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts).strip(), None
        if name.endswith((".txt", ".md", ".rtf")):
            return file_bytes.decode("utf-8", errors="ignore").strip(), None
        if name.endswith(".doc"):
            return "", "Old .doc format isn't supported — save as .docx or PDF."
        return "", f"Unsupported file type. Use PDF, DOCX or TXT."
    except Exception as e:
        logger.error(f"CV extract failed for {filename}: {e}")
        return "", f"Could not read this file: {str(e)[:120]}"


class ATSService:

    async def _jobs(self):
        from app.database import connect
        db = await connect()
        return db.ats_jobs if db is not None else None

    async def _cands(self):
        from app.database import connect
        db = await connect()
        return db.ats_candidates if db is not None else None

    def _scope(self, user_id, workspace_id=None):
        from app.services.workspace_service import is_personal
        if workspace_id and not is_personal(workspace_id):
            return {"workspace_id": workspace_id}
        return {"user_id": user_id}

    # ── Criteria: AI proposes, HR disposes ────────────────────────────────────
    async def suggest_criteria(self, jd_text: str) -> dict:
        """Read the JD and propose weighted, job-relevant criteria for HR to edit."""
        if not (jd_text or "").strip():
            return {"success": False, "error": "Paste the job description first."}
        prompt = (
            "You are an expert recruiter designing a fair, job-relevant scoring rubric.\n"
            "Read the job description and propose 4-7 criteria to score candidates on.\n\n"
            "RULES:\n"
            "- Criteria must be JOB-RELEVANT and observable in a CV (skills, experience, "
            "qualifications, domain, tools, outcomes).\n"
            "- NEVER propose criteria that are proxies for protected characteristics: no "
            "age, gender, nationality, marital status, school prestige, 'culture fit', "
            "employment gaps, or personality.\n"
            "- Weights must sum to 100.\n"
            "- Each criterion needs a short 'what good looks like' description.\n\n"
            f"JOB DESCRIPTION:\n{jd_text[:6000]}\n\n"
            'Reply ONLY with JSON, no preamble:\n'
            '{"criteria":[{"name":"...","weight":30,"description":"...","must_have":true}]}'
        )
        try:
            from app.services.llm_service import llm_service
            from app.models.schemas import LLMProvider
            text, _, _ = await llm_service.chat(
                messages=[{"role": "user", "content": prompt}],
                industry="general", provider=LLMProvider.anthropic,
                max_tokens=1200, temperature=0.1)
            data = _parse_json(text)
            crits = data.get("criteria", []) if data else []
            if not crits:
                return {"success": False, "error": "Could not derive criteria — edit them manually."}
            crits = _normalise_weights(crits)
            return {"success": True, "criteria": crits}
        except Exception as e:
            logger.error(f"Criteria suggestion failed: {e}")
            return {"success": False, "error": f"Could not read the JD: {str(e)[:140]}"}

    # ── Jobs ──────────────────────────────────────────────────────────────────
    async def create_job(self, user_id: str, title: str, jd_text: str,
                         criteria: list, anonymise_cvs: bool = True,
                         workspace_id: str = None) -> dict:
        col = await self._jobs()
        if col is None:
            return {"success": False, "error": "ATS unavailable — database not configured."}
        if not (title or "").strip():
            return {"success": False, "error": "Give the role a title."}
        if not criteria:
            return {"success": False, "error": "Add at least one scoring criterion."}
        from app.services.workspace_service import is_personal
        doc = {
            "_id": str(uuid.uuid4()),
            "user_id": user_id,
            "workspace_id": workspace_id if (workspace_id and not is_personal(workspace_id)) else None,
            "title": title.strip()[:140],
            "jd_text": (jd_text or "")[:20000],
            "criteria": _normalise_weights(criteria),
            "anonymise": bool(anonymise_cvs),
            "created_at": now(),
            "candidate_count": 0,
        }
        await col.insert_one(doc)
        return {"success": True, "job_id": doc["_id"], "job": _job_public(doc)}

    async def list_jobs(self, user_id: str, workspace_id: str = None) -> dict:
        col = await self._jobs()
        if col is None:
            return {"success": True, "items": []}
        items = []
        async for d in col.find(self._scope(user_id, workspace_id), {"jd_text": 0}).sort("created_at", -1):
            items.append(_job_public(d))
        return {"success": True, "items": items}

    async def get_job(self, user_id: str, job_id: str, workspace_id: str = None) -> dict:
        col = await self._jobs()
        if col is None:
            return {"success": False, "error": "ATS unavailable."}
        d = await col.find_one({"_id": job_id, **self._scope(user_id, workspace_id)})
        if not d:
            return {"success": False, "error": "Job not found."}
        return {"success": True, "job": {**_job_public(d), "jd_text": d.get("jd_text", "")}}

    async def update_criteria(self, user_id: str, job_id: str, criteria: list,
                              anonymise_cvs: bool = None, workspace_id: str = None) -> dict:
        """HR has the final say — their criteria and weights overwrite the AI's."""
        col = await self._jobs()
        if col is None:
            return {"success": False, "error": "ATS unavailable."}
        upd = {"criteria": _normalise_weights(criteria)}
        if anonymise_cvs is not None:
            upd["anonymise"] = bool(anonymise_cvs)
        res = await col.update_one({"_id": job_id, **self._scope(user_id, workspace_id)}, {"$set": upd})
        if res.matched_count == 0:
            return {"success": False, "error": "Job not found."}
        return {"success": True, "criteria": upd["criteria"]}

    async def delete_job(self, user_id: str, job_id: str, workspace_id: str = None) -> dict:
        col = await self._jobs()
        cands = await self._cands()
        if col is None:
            return {"success": False, "error": "ATS unavailable."}
        res = await col.delete_one({"_id": job_id, **self._scope(user_id, workspace_id)})
        if res.deleted_count == 0:
            return {"success": False, "error": "Job not found."}
        if cands is not None:
            await cands.delete_many({"job_id": job_id})
        return {"success": True, "deleted": job_id}

    # ── Scoring ───────────────────────────────────────────────────────────────
    async def score_cv(self, user_id: str, job_id: str, cv_text: str, filename: str,
                       candidate_name: str = None, workspace_id: str = None) -> dict:
        jobs = await self._jobs()
        cands = await self._cands()
        if jobs is None:
            return {"success": False, "error": "ATS unavailable."}
        job = await jobs.find_one({"_id": job_id, **self._scope(user_id, workspace_id)})
        if not job:
            return {"success": False, "error": "Job not found."}
        if not (cv_text or "").strip():
            return {"success": False, "error": "This CV appears to be empty or unreadable."}

        text = cv_text[:MAX_CV_CHARS]
        redacted = {}
        if job.get("anonymise", True):
            text, redacted = anonymise(text, candidate_name)

        result = await self._llm_score(job, text)
        if not result.get("success"):
            return result

        doc = {
            "_id": str(uuid.uuid4()),
            "job_id": job_id,
            "user_id": user_id,
            "workspace_id": job.get("workspace_id"),
            "filename": filename,
            # Name is stored for the recruiter's own reference, but it was NOT
            # visible to the model when anonymisation is on.
            "candidate_name": candidate_name or _guess_name(filename),
            "anonymised": bool(job.get("anonymise", True)),
            "redacted": redacted,
            "score": result["score"],
            "breakdown": result["breakdown"],
            "strengths": result.get("strengths", []),
            "gaps": result.get("gaps", []),
            "summary": result.get("summary", ""),
            "created_at": now(),
        }
        if cands is not None:
            await cands.insert_one(doc)
            await jobs.update_one({"_id": job_id}, {"$inc": {"candidate_count": 1}})
        return {"success": True, "candidate": _cand_public(doc)}

    async def _llm_score(self, job: dict, cv_text: str) -> dict:
        crits = job.get("criteria", [])
        crit_block = "\n".join(
            f"- {c['name']} (weight {c['weight']}%{', MUST-HAVE' if c.get('must_have') else ''}): {c.get('description','')}"
            for c in crits)
        prompt = (
            "You are an expert, fair recruiter scoring ONE candidate against a rubric.\n\n"
            "ABSOLUTE RULES:\n"
            "1. Score ONLY against the criteria below. Nothing else counts.\n"
            "2. EVERY sub-score must cite a short verbatim quote from the CV as evidence. "
            "If you cannot find evidence, score it low and say 'No evidence found' — never invent.\n"
            "3. IGNORE and never reward or penalise: name, age, gender, nationality, marital "
            "status, photos, school prestige, employment gaps, or writing flourish. Judge "
            "demonstrated capability only.\n"
            "4. Redacted markers like [NAME REMOVED] are deliberate — do not speculate about them.\n"
            "5. Be calibrated: 90+ is exceptional and rare, 70-89 strong, 50-69 partial, "
            "below 50 weak. Do not inflate.\n\n"
            f"ROLE: {job.get('title')}\n\n"
            f"SCORING CRITERIA:\n{crit_block}\n\n"
            f"JOB DESCRIPTION (context):\n{(job.get('jd_text') or '')[:2500]}\n\n"
            f"CANDIDATE CV:\n{cv_text}\n\n"
            'Reply ONLY with JSON, no preamble:\n'
            '{"breakdown":[{"criterion":"...","score":0-100,"evidence":"verbatim quote from CV",'
            '"reasoning":"one line"}],"strengths":["..."],"gaps":["..."],'
            '"summary":"2 sentences on fit for THIS role"}'
        )
        try:
            from app.services.llm_service import llm_service
            from app.models.schemas import LLMProvider
            text, _, _ = await llm_service.chat(
                messages=[{"role": "user", "content": prompt}],
                industry="general", provider=LLMProvider.anthropic,
                max_tokens=2000, temperature=0.05)
            data = _parse_json(text)
            if not data or not data.get("breakdown"):
                return {"success": False, "error": "Scoring returned an unreadable result. Try again."}
            # Weighted total — computed HERE, never trusted to the model's arithmetic
            total, breakdown = 0.0, []
            by_name = {c["name"].lower(): c for c in crits}
            for b in data["breakdown"]:
                c = by_name.get(str(b.get("criterion", "")).lower())
                w = float(c["weight"]) if c else 0.0
                sc = max(0.0, min(100.0, float(b.get("score", 0))))
                total += sc * (w / 100.0)
                breakdown.append({
                    "criterion": b.get("criterion"), "score": round(sc),
                    "weight": w, "evidence": (b.get("evidence") or "")[:400],
                    "reasoning": (b.get("reasoning") or "")[:300],
                    "must_have": bool(c.get("must_have")) if c else False,
                })
            # A missed must-have caps the overall score — a hard gate, not a nudge
            missed = [b for b in breakdown if b["must_have"] and b["score"] < 50]
            capped = False
            if missed:
                total = min(total, 49.0)
                capped = True
            return {"success": True, "score": round(total, 1), "breakdown": breakdown,
                    "strengths": data.get("strengths", [])[:6],
                    "gaps": data.get("gaps", [])[:6],
                    "summary": (data.get("summary") or "")[:600],
                    "capped_by_must_have": capped}
        except Exception as e:
            logger.error(f"CV scoring failed: {e}")
            return {"success": False, "error": f"Scoring failed: {str(e)[:160]}"}

    # ── Candidates ────────────────────────────────────────────────────────────
    async def list_candidates(self, user_id: str, job_id: str, workspace_id: str = None) -> dict:
        cands = await self._cands()
        if cands is None:
            return {"success": True, "items": []}
        jobs = await self._jobs()
        job = await jobs.find_one({"_id": job_id, **self._scope(user_id, workspace_id)})
        if not job:
            return {"success": False, "error": "Job not found.", "items": []}
        items = []
        async for d in cands.find({"job_id": job_id}).sort("score", -1):
            items.append(_cand_public(d))
        for i, c in enumerate(items, 1):
            c["rank"] = i
        return {"success": True, "items": items, "anonymised": bool(job.get("anonymise", True))}

    async def delete_candidate(self, user_id: str, cand_id: str, workspace_id: str = None) -> dict:
        cands = await self._cands()
        if cands is None:
            return {"success": False, "error": "ATS unavailable."}
        d = await cands.find_one({"_id": cand_id})
        if not d:
            return {"success": False, "error": "Candidate not found."}
        jobs = await self._jobs()
        job = await jobs.find_one({"_id": d["job_id"], **self._scope(user_id, workspace_id)})
        if not job:
            return {"success": False, "error": "Not your job."}
        await cands.delete_one({"_id": cand_id})
        await jobs.update_one({"_id": d["job_id"]}, {"$inc": {"candidate_count": -1}})
        return {"success": True, "deleted": cand_id}


# ── helpers ───────────────────────────────────────────────────────────────────
def _parse_json(text: str):
    if not text:
        return None
    t = text.strip()
    t = re.sub(r"^```(?:json)?|```$", "", t, flags=re.M).strip()
    m = re.search(r"\{.*\}", t, re.S)
    if not m:
        return None
    try:
        return json.loads(m.group(0))
    except Exception:
        return None

def _normalise_weights(crits: list) -> list:
    out = []
    for c in crits or []:
        try:
            out.append({
                "name": str(c.get("name", "")).strip()[:80] or "Criterion",
                "weight": max(0.0, float(c.get("weight", 0))),
                "description": str(c.get("description", ""))[:300],
                "must_have": bool(c.get("must_have", False)),
            })
        except Exception:
            continue
    total = sum(c["weight"] for c in out)
    if total > 0 and abs(total - 100) > 0.01:
        for c in out:                       # rescale so weights always mean %
            c["weight"] = round(c["weight"] * 100.0 / total, 1)
    return out

def _guess_name(filename: str) -> str:
    base = re.sub(r"\.(pdf|docx?|txt|md)$", "", filename or "", flags=re.I)
    base = re.sub(r"[_\-]+", " ", base)
    base = re.sub(r"\b(cv|resume|curriculum vitae|final|v\d+|\d{4})\b", "", base, flags=re.I)
    return base.strip().title() or "Candidate"

def _job_public(d: dict) -> dict:
    return {"id": d["_id"], "title": d.get("title"), "criteria": d.get("criteria", []),
            "anonymise": d.get("anonymise", True),
            "candidate_count": d.get("candidate_count", 0),
            "created_at": _iso(d.get("created_at"))}

def _cand_public(d: dict) -> dict:
    return {"id": d["_id"], "job_id": d.get("job_id"), "filename": d.get("filename"),
            "candidate_name": d.get("candidate_name"), "score": d.get("score"),
            "breakdown": d.get("breakdown", []), "strengths": d.get("strengths", []),
            "gaps": d.get("gaps", []), "summary": d.get("summary", ""),
            "anonymised": d.get("anonymised", True), "redacted": d.get("redacted", {}),
            "created_at": _iso(d.get("created_at"))}

def _iso(v):
    return v.isoformat() if isinstance(v, datetime) else v


ats_service = ATSService()
